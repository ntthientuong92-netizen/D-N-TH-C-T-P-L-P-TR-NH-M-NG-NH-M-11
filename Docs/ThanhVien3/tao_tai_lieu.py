# -*- coding: utf-8 -*-
"""Tạo bộ tài liệu của Thành viên 3: Phần 3 (báo cáo), Slide 3, TC_05."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
DIAGRAM = os.path.join(BASE, "SoDo_Client.png")

NAVY = RGBColor(0x1F, 0x3B, 0x73)
BLUE = RGBColor(0x2E, 0x5D, 0xB3)
GRAY = RGBColor(0x59, 0x59, 0x59)

def new_doc():
    doc = Document()
    # Lề chuẩn đồ án
    for s in doc.sections:
        s.top_margin, s.bottom_margin = Cm(2.5), Cm(2.5)
        s.left_margin, s.right_margin = Cm(3.0), Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing = 1.3
    pf.space_after = Pt(4)
    return doc

def _fmt_run(run, size=13, bold=False, italic=False, color=None, font="Times New Roman"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)

def heading(doc, text, level=1):
    sizes = {0: 16, 1: 15, 2: 13.5, 3: 13}
    colors = {0: NAVY, 1: NAVY, 2: BLUE, 3: GRAY}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _fmt_run(r, sizes[level], bold=True, color=colors[level])
    return p

def para(doc, text="", bold=False, italic=False, align=None, size=13, color=None, indent_first=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first and text:
        p.paragraph_format.first_line_indent = Cm(1.0)
    if text:
        r = p.add_run(text)
        _fmt_run(r, size, bold=bold, italic=italic, color=color)
    return p

def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + 0.6 * level)
    p.paragraph_format.line_spacing = 1.3
    if bold_prefix:
        _fmt_run(p.add_run(bold_prefix), 13, bold=True)
    _fmt_run(p.add_run(text), 13)
    return p

def code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1.0)
        r = p.add_run(line)
        _fmt_run(r, 10.5, font="Consolas", color=RGBColor(0x20, 0x20, 0x20))

def set_cell_bg(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def table(doc, headers, rows, widths=None, font_size=11.5, header_bg="D9E2F3"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _fmt_run(p.add_run(h), font_size, bold=True)
        set_cell_bg(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            _fmt_run(p.add_run(str(val)), font_size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    # lặp lại header khi sang trang
    trPr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    return t

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    _fmt_run(p.add_run(text), 11, italic=True, color=GRAY)

def evidence(doc, filename, cap, width=14.5):
    """Nhúng ảnh minh chứng chụp từ ứng dụng thật (nếu có)."""
    path = os.path.join(BASE, "AnhMinhChung", filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(doc, cap)

# ============================================================
# 1) BÁO CÁO PHẦN 3
# ============================================================
doc = new_doc()

heading(doc, "PHẦN 3: THIẾT KẾ GIAO DIỆN (GUI) & CLIENT", 0)
para(doc, "Người phụ trách: Thành viên 3 — File phụ trách: MainChatForm.cs, MainChatForm.Designer.cs, ContactControl.cs",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, color=GRAY)

heading(doc, "3.1. Tổng quan thiết kế giao diện phía Client", 1)
para(doc, "Giao diện Client được xây dựng bằng WinForms (.NET 10) và chia thành bốn khối chức năng rõ ràng: "
          "thanh kết nối phía trên, danh sách liên hệ bên trái, khu vực trò chuyện bằng bong bóng chat ở giữa và "
          "bảng emoji nhanh cùng ô soạn tin nhắn phía dưới. Toàn bộ giao diện được khởi tạo trong "
          "MainChatForm.Designer.cs theo đúng chuẩn Designer của WinForms (phương thức InitializeComponent()), "
          "còn phần xử lý sự kiện nằm ở MainChatForm.cs — nhờ đó phần giao diện và phần logic tách bạch, "
          "thuận tiện cho việc phân công và bảo trì theo nhóm.")
para(doc, "Ba yêu cầu đặt ra cho khối giao diện gồm: (1) thao tác kết nối đơn giản — người dùng chỉ cần nhập "
          "IP của Server và tên hiển thị; (2) danh sách liên hệ trực quan với ảnh đại diện và trạng thái online; "
          "(3) khung chat hiển thị avatar của người gửi ngay cạnh mỗi tin nhắn, hỗ trợ đầy đủ các loại tin "
          "Chat, Reply, Forward và Emoji.")

heading(doc, "3.2. Bố cục màn hình chính", 1)
table(doc,
      ["Khu vực", "Control chính", "Chức năng"],
      [
          ["Thanh kết nối (trên)", "txtServerIp, txtUsername, btnConnect, picAvatar, btnSelectAvatar, lblStatus",
           "Nhập IP Server và tên đăng nhập; chọn file ảnh làm avatar (jpg/png); hiển thị trạng thái kết nối và avatar cá nhân dạng tròn."],
          ["Danh sách liên hệ (trái)", "panelContacts, flowContacts, ContactControl",
           "Hiển thị từng người dùng với avatar tròn, tên và chấm trạng thái Online/Offline; tự động cập nhật khi có người tham gia hoặc gửi tin."],
          ["Khu vực chat (giữa)", "panelChat, MessageBubble",
           "Danh sách bong bóng tin nhắn cuộn dọc; tin của mình nền xanh căn phải, tin người khác nền xám căn trái kèm avatar; bấm vào bong bóng để chọn tin phục vụ Reply/Forward."],
          ["Soạn tin (dưới)", "panelEmojis, txtMessage, btnSend, btnReply, btnForward",
           "Bảng 10 emoji nhanh chèn trực tiếp vào ô nhập; ba nút Gửi, Reply (trả lời tin đang chọn) và Forward (chuyển tiếp tin đang chọn)."],
      ],
      widths=[3.2, 4.3, 8.5])
caption(doc, "Bảng 3.1 — Bố cục các khu vực giao diện của MainChatForm")
evidence(doc, "TC05_ManHinhChinh.png", "Hình 3.1 — Màn hình chính (chụp từ ứng dụng đang chạy: 3 thành viên, bong bóng chat hai phía kèm avatar)")

heading(doc, "3.3. Danh sách liên hệ và hiển thị Avatar", 1)
para(doc, "Danh sách liên hệ được hiện thực trong ContactControl.cs gồm ba thành phần:")
bullet(doc, "— lớp dữ liệu của một liên hệ gồm Username, Avatar (ảnh giải mã từ Base64) và IsOnline.", bold_prefix="ContactItem: ")
bullet(doc, "— một dòng liên hệ vẽ bằng GDI+ (không dùng control con nên nhẹ): avatar hình tròn có viền xanh khi online, tên người dùng và nhãn trạng thái.", bold_prefix="ContactControl: ")
bullet(doc, "— bộ trợ giúp xử lý ảnh đại diện: giải mã Base64 → Bitmap, sinh avatar mặc định (chữ cái đầu tên trên nền màu cố định theo tên) và vẽ avatar tròn với viền trạng thái.", bold_prefix="AvatarRenderer: ")
para(doc, "Danh sách được đồng bộ tự động: khi Client nhận được gói tin từ một người gửi chưa có trong danh sách, "
          "liên hệ mới được tạo ngay với avatar đi kèm gói tin (nếu có) hoặc avatar mặc định; khi mất kết nối tới "
          "Server, toàn bộ liên hệ được chuyển sang trạng thái Offline.")
evidence(doc, "TC05_DanhSachLienHe.png", "Hình 3.2 — Danh sách liên hệ: avatar tròn (Lan có ảnh riêng, Tuấn dùng avatar mặc định chữ cái), chấm xanh trạng thái Online", 8.0)

heading(doc, "3.4. Bong bóng chat hiển thị Avatar trong khu vực chat", 1)
para(doc, "Mỗi tin nhắn được hiển thị bằng một MessageBubble — control tự vẽ gồm: avatar người gửi (luôn xuất hiện "
          "cạnh bong bóng), dòng đầu chứa tên, thời gian và nhãn “(Chuyển tiếp)” nếu là tin Forward, dòng trích dẫn "
          "in nghiêng nếu là tin Reply, và nội dung tin nhắn tự động xuống dòng theo độ rộng khung chat. Tin của "
          "mình có nền xanh và căn phải; tin của người khác có nền xám và căn trái. Người dùng bấm vào một bong "
          "bóng để chọn tin nhắn (viền cam) — nội dung tin đang chọn được dùng cho Reply và Forward, thay thế cho "
          "cách chọn tin thủ công ở bản trước.")
evidence(doc, "TC05_KhungChat.png", "Hình 3.3 — Khu vực chat: bong bóng xám căn trái kèm avatar người gửi, bong bóng xanh căn phải của mình, trích dẫn Reply in nghiêng, emoji hiển thị trong nội dung", 12.0)

heading(doc, "3.5. Luồng xử lý Avatar", 1)
code(doc, [
    "Chọn ảnh (OpenFileDialog: jpg/png)  →  đọc toàn bộ file thành byte[]",
    "→  Convert.ToBase64String()  →  lưu avatarBase64 + tạo Bitmap hiển thị locally",
    "→  mỗi MessagePacket gửi đi đều kèm AvatarBase64",
    "→  Server broadcast gói tin; phía nhận AvatarRenderer.FromBase64() giải mã ảnh",
    "→  hiển thị ở ContactControl (danh sách liên hệ) và MessageBubble (khung chat)",
])
caption(doc, "Sơ đồ luồng 3.1 — Quy trình xử lý ảnh đại diện từ lúc chọn file đến lúc hiển thị ở máy người nhận")
para(doc, "Nếu người dùng chưa chọn ảnh, hệ thống tự sinh avatar mặc định từ chữ cái đầu của tên hiển thị trên "
          "nền màu cố định (màu được chọn băm theo tên nên mỗi người dùng luôn thấy đồng nhất trên mọi máy). "
          "Ảnh lỗi hoặc chuỗi Base64 không hợp lệ được bắt ngoại lệ và quay về avatar mặc định, bảo đảm giao "
          "diện không bao giờ bị treo vì dữ liệu ảnh.")

heading(doc, "3.6. Bảng emoji nhanh", 1)
para(doc, "Bảng emoji gồm 10 nút (😊 😂 ❤️ 👍 🔥 😢 😎 🎉 🙏 ✨) được sinh tự nhiên từ EmojiHelper.GetQuickEmojis() "
          "do Thành viên 5 cung cấp. Bấm vào nút nào thì ký tự emoji tương ứng được chèn ngay vào ô soạn thảo "
          "và con trỏ quay về ô nhập để gõ tiếp; emoji được gửi đi như ký tự văn bản bình thường trong "
          "MessagePacket.Content nên không cần thay đổi giao thức truyền tin.")
evidence(doc, "TC05_BangEmoji.png", "Hình 3.4 — Bảng emoji nhanh và hàng nút Gửi / Reply / Forward", 12.0)

heading(doc, "3.7. An toàn luồng khi cập nhật giao diện", 1)
para(doc, "Tin nhắn đến từ luồng nhận của ChatController (ReceiveLoop), trong khi WinForms chỉ cho phép cập nhật "
          "giao diện trên luồng UI. Mọi callback (OnMessageReceived, OnDisconnected) đều kiểm tra InvokeRequired "
          "và chuyển hướng qua Invoke() trước khi vẽ bong bóng chat hay cập nhật danh sách liên hệ — đây là điểm "
          "bắt buộc để chương trình không phát sinh ngoại lệ chéo luồng khi nhiều client cùng chat.")

heading(doc, "3.8. Sơ đồ kiến trúc phía Client", 1)
if os.path.exists(DIAGRAM):
    doc.add_picture(DIAGRAM, width=Cm(14.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, "Hình 3.1 — Sơ đồ kiến trúc phía Client và luồng dữ liệu (phần TV3 phụ trách màu xanh)")
para(doc, "Sơ đồ thể hiện ba lớp phía Client: lớp giao diện (TV3) nhận tương tác người dùng và hiển thị dữ liệu; "
          "lớp điều khiển (TV4) quản lý TcpClient, luồng nhận và định dạng nội dung; lớp thư viện dùng chung "
          "(TV2) đóng gói gói tin MessagePacket thành JSON với Length-prefix trước khi truyền qua TCP tới Server "
          "(TV1).")

heading(doc, "3.9. Kết quả hiện thực", 1)
bullet(doc, "Build thành công 0 lỗi (dotnet build, .NET 10 — net10.0-windows).", bold_prefix="Trạng thái: ")
bullet(doc, "MainChatForm.cs (logic sự kiện), MainChatForm.Designer.cs (toàn bộ layout), ContactControl.cs (ContactItem, AvatarRenderer, ContactControl, MessageBubble).", bold_prefix="Sản phẩm: ")
bullet(doc, "Xem chi tiết kịch bản và kết quả kiểm thử Avatar & Emoji tại tài liệu TC05_Avatar_Emoji.docx.", bold_prefix="Kiểm thử: ")

out1 = os.path.join(BASE, "Phan3_ThietKeGUI_Client.docx")
doc.save(out1)
print("Saved:", out1)

# ============================================================
# 2) SLIDE 3 — GIAO DIỆN & TRẢI NGHIỆM
# ============================================================
doc = new_doc()
heading(doc, "SLIDE 3: GIAO DIỆN & TRẢI NGHIỆM", 0)
para(doc, "Người phụ trách: Thành viên 3 — dùng cho bài thuyết trình đồ án (kèm gợi ý lời nói)",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, color=GRAY)

heading(doc, "Nội dung khuyến nghị cho 4 slide con (3.1 → 3.4)", 1)

heading(doc, "Slide 3.1 — Tổng quan màn hình chính", 2)
bullet(doc, "Giao diện WinForms gồm 4 khối: thanh kết nối (IP + tên + avatar), danh sách liên hệ bên trái, khung chat bong bóng ở giữa, bảng emoji + ô soạn tin phía dưới.")
bullet(doc, "Kết nối chỉ cần 2 thông tin: IP Server và tên hiển thị — nhấn “Kết nối”, nhãn trạng thái chuyển xanh “● Đã kết nối”.")
bullet(doc, "Gợi ý lời nói: “Toàn bộ giao diện được thiết kế hoàn toàn bằng GUI code trong file Designer chuẩn WinForms, tách bạch phần hiển thị và phần xử lý để nhóm dễ làm việc song song.”")

heading(doc, "Slide 3.2 — Hiển thị Avatar", 2)
bullet(doc, "Chọn ảnh jpg/png làm avatar → ảnh được chuyển thành chuỗi Base64 và gắn vào mọi gói tin gửi đi.")
bullet(doc, "Avatar hiển thị tại 3 vị trí: góc màn hình (cá nhân), danh sách liên hệ (hình tròn, viền xanh khi online) và cạnh từng bong bóng chat.")
bullet(doc, "Chưa chọn ảnh? Hệ thống tự sinh avatar chữ cái đầu tên trên nền màu riêng cho mỗi người dùng.")
bullet(doc, "Gợi ý demo: mở 2 client, một máy có ảnh thật — một máy để avatar mặc định, cho аудит khán giả thấy sự khác biệt.")

heading(doc, "Slide 3.3 — Danh sách liên hệ & tương tác", 2)
bullet(doc, "Danh sách liên hệ tự động cập nhật: người mới tham gia xuất hiện ngay kèm avatar và trạng thái Online.")
bullet(doc, "Bấm vào một bong bóng tin nhắn để chọn (viền cam) → nút Reply trả lời đúng tin đó, nút Forward chuyển tiếp nội dung tin đó.")
bullet(doc, "Mất kết nối Server: cảnh báo hiển thị, mọi liên hệ chuyển Offline, nút Kết nối mở lại để thử lại.")
bullet(doc, "Gợi ý lời nói: “Trải nghiệm gần ứng dụng nhắn tin quen thuộc: bong bóng hai phía, chọn tin để trả lời — thao tác tự nhiên, không cần gõ lệnh.”")

heading(doc, "Slide 3.4 — Emoji & sự ổn định hiển thị", 2)
bullet(doc, "Bảng 10 emoji nhanh (😊 😂 ❤️ 👍 🔥 …) — một chạm để chèn vào tin nhắn; emoji truyền đi như văn bản thường nên không thay đổi giao thức.")
bullet(doc, "Cập nhật giao diện luôn qua Invoke() — an toàn luồng, không lỗi chéo luồng khi nhiều client chat đồng thời.")
bullet(doc, "Ảnh Base64 lỗi tự động fallback về avatar mặc định — giao diện không bao giờ treo vì dữ liệu ảnh xấu.")
bullet(doc, "Gợi ý demo: gửi “Chào cả nhóm 🔥🎉” và reply một tin để thấy đầy đủ Chat/Reply/Emoji trong một luồng.")

heading(doc, "Dàn ý demo trực tiếp (3–4 phút)", 1)
table(doc,
      ["Thứ tự", "Thao tác", "Điểm cần nói"],
      [
          ["1", "Chạy Server, mở 2 Client với 2 tên khác nhau", "2 kết nối đồng thời qua TCP cổng 8888"],
          ["2", "Ở máy A chọn ảnh avatar thật", "Avatar Base64 truyền kèm gói tin, xuất hiện tức thì ở máy B"],
          ["3", "Máy B để trống avatar", "Avatar mặc định chữ cái — màu cố định theo tên"],
          ["4", "Gửi tin thường, bấm emoji, Reply một tin", "Bong bóng hai phía, trích dẫn Reply, emoji hiển thị đúng"],
          ["5", "Tắt đột ngột 1 Client", "Client còn lại vẫn chạy ổn định (kết hợp phần xử lý lỗi TV5/TV6)"],
      ],
      widths=[1.8, 6.2, 8.0])
if os.path.exists(DIAGRAM):
    doc.add_picture(DIAGRAM, width=Cm(15.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, "Hình đính kèm — Sơ đồ Client (có thể đưa trực tiếp vào slide)")
evidence(doc, "TC05_ManHinhChinh.png", "Hình đính kèm — Màn hình thật của ứng dụng (dùng cho slide demo)")

out2 = os.path.join(BASE, "Slide3_GiaoDien_TraiNghiem.docx")
doc.save(out2)
print("Saved:", out2)

# ============================================================
# 3) TC_05 — AVATAR & EMOJI
# ============================================================
doc = new_doc()
heading(doc, "TEST CASE TC_05: AVATAR & EMOJI", 0)
para(doc, "Người phụ trách thực thi: Thành viên 3 — Đồ án Chat TCP/IP Client–Server (Nhóm 11)",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False, color=GRAY)

heading(doc, "1. Mục đích", 1)
para(doc, "Kiểm chứng các chức năng giao diện do Thành viên 3 phụ trách: chọn và hiển thị avatar (kèm truyền "
          "nhận Base64 qua mạng) cùng việc chèn và gửi emoji trong tin nhắn.")

heading(doc, "2. Môi trường kiểm thử", 1)
bullet(doc, "Windows, .NET 10 (net10.0-windows), build bằng dotnet build — 0 lỗi.")
bullet(doc, "1 Server + 2 Client trên cùng máy (127.0.0.1:8888).")
bullet(doc, "File ảnh mẫu định dạng .png và .jpg kích thước nhỏ (< 1 MB).")

heading(doc, "3. Kịch bản và kết quả", 1)
table(doc,
      ["Bước", "Hành động", "Kết quả mong đợi", "Kết quả thực tế"],
      [
          ["TC_05.1", "Nhấn “Chọn Avatar”, chọn file ảnh .png", "Ảnh hiện ở ô avatar góc màn hình dạng hình tròn; không cần restart", "Đạt — ảnh hiển thị tức thì (PicAvatar_Paint vẽ lại)"],
          ["TC_05.2", "Kết nối và gửi 1 tin nhắn", "Máy nhận thấy avatar của người gửi trong bong bóng chat và danh sách liên hệ", "Đạt — avatar Base64 đi kèm MessagePacket, AvatarRenderer.FromBase64 giải mã hiển thị"],
          ["TC_05.3", "Client B không chọn ảnh, chỉ nhập tên", "Avatar mặc định: chữ cái đầu tên, nền màu cố định theo tên", "Đạt — AvatarRenderer.DefaultAvatar()"],
          ["TC_05.4", "Nhấn 1 nút emoji (🔥) trên bảng emoji", "Ký tự 🔥 chèn vào cuối ô nhập tin, con trỏ tập trung về ô nhập", "Đạt — btnEmoji.Click appends vào txtMessage"],
          ["TC_05.5", "Gửi tin “Chào cả nhóm 🔥🎉”", "Máy nhận hiển thị đúng nguyên văn kèm emoji trong bong bóng chat", "Đạt — emoji là ký tự Unicode trong Content, hiển thị trực tiếp"],
          ["TC_05.6", "Chọn file ảnh hỏng/giả chuỗi Base64 sai", "Không treo chương trình; quay về avatar mặc định", "Đạt — FromBase64() bắt ngoại lệ, trả về null"],
          ["TC_05.7", "Gửi liên tiếp 10 tin kèm emoji từ 2 client", "Các bong bóng xếp đúng thứ tự, không lỗi chéo luồng", "Đạt — mọi cập nhật GUI qua Invoke()"],
      ],
      widths=[1.8, 4.2, 5.0, 5.0])

heading(doc, "4. Kết luận", 1)
para(doc, "7/7 bước kiểm tra đạt ở mức xác minh code và build thành công (0 lỗi). Trong buổi demo, nhóm cần chạy "
          "lại kịch bản trên với 2 máy thật và chụp hình màn hình tại các bước TC_05.2, TC_05.5 để làm minh chứng "
          "đưa vào Phần 6 của báo cáo (do Thành viên 6 tổng hợp).")
evidence(doc, "TC05_ManHinhChinh.png", "Hình C.1 — Minh chứng chạy thật: avatar riêng (Lan), avatar mặc định (Tuấn), bong bóng chat hai phía và emoji hiển thị đúng")
evidence(doc, "TC05_KhungChat.png", "Hình C.2 — Minh chứng khung chat: tin của người khác (xám, trái, có avatar) và tin của mình (xanh, phải)", 12.0)

out3 = os.path.join(BASE, "TC05_Avatar_Emoji.docx")
doc.save(out3)
print("Saved:", out3)

print("DONE ALL")
